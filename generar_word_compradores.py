#!/usr/bin/env python3
"""Script para generar documentos Word individuales para compradores"""

import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def generar_documentos_word_compradores():
    """Función para generar documentos Word para compradores"""
    
    # Cargar datos desde el JSON generado
    with open("agenda_completa.json", 'r', encoding='utf-8') as f:
        resultado = json.load(f)
    
    # Crear carpeta para documentos si no existe
    carpeta_docs = "documentos_compradores"
    if not os.path.exists(carpeta_docs):
        os.makedirs(carpeta_docs)
        print(f"Carpeta creada: {carpeta_docs}")

    compradores_procesados = 0
    
    # Generar un documento para cada comprador
    for comprador, datos_comprador in resultado["resumen_por_comprador"].items():
        if datos_comprador["total_citas"] > 0:  # Solo generar para compradores con citas
            print(f"Procesando comprador: {comprador}")
            
            doc = Document()
            
            # Título principal: nombre del comprador
            titulo = doc.add_heading(comprador, 0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Agregar espacio
            doc.add_paragraph("")
            
            # Información general
            info_p = doc.add_paragraph()
            info_p.add_run("Total de citas programadas: ").bold = True
            info_p.add_run(str(datos_comprador["total_citas"]))
            
            # Agregar espacio
            doc.add_paragraph("")
            
            # Crear tabla con citas
            # Encabezados: Franja Horaria | Mesa | Vendedores
            tabla = doc.add_table(rows=1, cols=3)
            tabla.style = 'Table Grid'
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Encabezados de la tabla
            encabezados = tabla.rows[0].cells
            encabezados[0].text = "FRANJA HORARIA"
            encabezados[1].text = "MESA"
            encabezados[2].text = "VENDEDORES"
            
            # Formatear encabezados
            for celda in encabezados:
                for parrafo in celda.paragraphs:
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in parrafo.runs:
                        run.font.bold = True
            
            # Agregar las citas del comprador
            for cita in datos_comprador["citas"]:
                fila = tabla.add_row().cells
                fila[0].text = cita["horario"]
                fila[1].text = ""  # Mesa vacía para completar manualmente
                
                # Unir vendedores con separador
                vendedores_texto = " | ".join(cita["vendedores"])
                fila[2].text = vendedores_texto
                
                # Centrar el contenido de las celdas
                for celda in fila:
                    for parrafo in celda.paragraphs:
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ajustar ancho de columnas
            for fila in tabla.rows:
                fila.cells[0].width = Inches(2.5)  # Franja horaria
                fila.cells[1].width = Inches(1.5)  # Mesa
                fila.cells[2].width = Inches(4.0)  # Vendedores (más ancho para múltiples vendedores)
            
            # Agregar espacio al final
            doc.add_paragraph("")
            
            # Nota al pie
            nota = doc.add_paragraph()
            nota.add_run("Nota: ").bold = True
            nota.add_run("Complete la columna 'Mesa' según la asignación del evento. Los vendedores están separados por '|' cuando hay múltiples.")
            nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Guardar el documento
            # Limpiar caracteres problemáticos del nombre del archivo
            caracteres_problematicos = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
            nombre_limpio = comprador
            for char in caracteres_problematicos:
                nombre_limpio = nombre_limpio.replace(char, '_')
            
            nombre_archivo = f"{nombre_limpio}.docx"
            ruta_completa = os.path.join(carpeta_docs, nombre_archivo)
            
            try:
                doc.save(ruta_completa)
                compradores_procesados += 1
                print(f"✓ Documento generado: {nombre_archivo}")
            except Exception as e:
                print(f"✗ Error al guardar documento para {comprador}: {e}")
    
    print(f"\n📄 DOCUMENTOS WORD GENERADOS PARA COMPRADORES:")
    print(f"   • Total documentos: {compradores_procesados}")
    print(f"   • Carpeta: {carpeta_docs}")
    print(f"   • Cada documento contiene: Nombre del comprador, tabla con franja horaria, mesa (vacía) y vendedores")
    
    return compradores_procesados

if __name__ == "__main__":
    print("🚀 Iniciando generación de documentos Word para compradores...")
    generar_documentos_word_compradores()
    print("✅ Proceso completado!")