#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Programa para organizar agenda de rueda de negocios
25 vendedores, 10 compradores
Tiempo: 10:15 AM - 1:00 PM (2h 45min)
Citas de 15 minutos, máximo 11 por vendedor (sin restricción práctica)
Formato: Distribución estratégica por franja: 5 citas (2 vendedores) + 5 citas (3 vendedores) = 25 vendedores activos
Todos los compradores participan en cada slot de tiempo
No se repiten encuentros comprador-vendedor
Múltiples citas simultáneas permitidas
"""

import random
import csv
import os
from datetime import datetime, timedelta
from typing import List, Dict, Tuple, Optional, Set
import json

class AgendaRuedaNegocios:
    def __init__(self):
        self.inicio = datetime.strptime("08:30", "%H:%M")  # Extendido para Regional SAS
        self.fin = datetime.strptime("13:00", "%H:%M")
        self.duracion_cita = 15  # minutos - vuelto a 15 minutos
        self.max_citas_vendedor = 11  # Sin restricción práctica - pueden tener hasta 11 citas
        self.num_vendedores = 25
        self.num_compradores = 10
        self.vendedores_por_cita = 3  # Cambiado: 3 vendedores por cita por defecto (nueva restricción)
        
        # Calcular slots de tiempo disponibles
        total_minutos = int((self.fin - self.inicio).total_seconds() / 60)
        self.num_slots = total_minutos // self.duracion_cita
        
        # Generar horarios
        self.horarios = self._generar_horarios()
        
        # Inicializar listas vacías - se llenarán al cargar preferencias
        self.vendedores = []
        self.compradores = []
        
        # Agenda: {slot: [(comprador, [vendedor1, vendedor2, vendedor3]), ...]}
        self.agenda = {i: [] for i in range(self.num_slots)}
        
        # Tracking para evitar repeticiones comprador-vendedor
        self.encuentros_realizados = {}  # {comprador: set(vendedores_ya_vistos)}
        
        # Contadores - se inicializarán después de cargar participantes
        self.citas_por_vendedor = {}
        self.citas_por_comprador = {}
        
        # Preferencias de citas (vendedor -> lista de compradores preferidos)
        self.preferencias_citas = {}
        
        # Citas ya asignadas para cumplir preferencias
        self.citas_preferencia_asignadas = set()
        
        # Flag para saber si ya se cargaron los participantes
        self.participantes_cargados = False

    def _generar_horarios(self) -> List[str]:
        """Genera la lista de horarios disponibles"""
        horarios = []
        tiempo_actual = self.inicio
        
        for i in range(self.num_slots):
            fin_slot = tiempo_actual + timedelta(minutes=self.duracion_cita)
            horario = f"{tiempo_actual.strftime('%H:%M')} - {fin_slot.strftime('%H:%M')}"
            horarios.append(horario)
            tiempo_actual = fin_slot
            
        return horarios

    def cargar_preferencias_archivo(self, ruta_archivo: str) -> bool:
        """Carga las preferencias de citas desde un archivo CSV"""
        try:
            if not os.path.exists(ruta_archivo):
                print(f"Archivo {ruta_archivo} no encontrado.")
                return False
            
            # Conjuntos temporales para recopilar nombres únicos
            vendedores_set = set()
            compradores_set = set()
            
            with open(ruta_archivo, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                
                # Saltar encabezado si existe
                primera_fila = next(reader, None)
                if primera_fila and ('vendedor' in primera_fila[0].lower() or 'nombre' in primera_fila[0].lower()):
                    pass  # Era encabezado, continuar con los datos
                else:
                    # No era encabezado, procesarla como dato
                    if primera_fila and len(primera_fila) >= 2:
                        vendedor = primera_fila[0].strip()
                        comprador = primera_fila[1].strip()
                        if vendedor and comprador:
                            vendedores_set.add(vendedor)
                            compradores_set.add(comprador)
                            if vendedor not in self.preferencias_citas:
                                self.preferencias_citas[vendedor] = []
                            if comprador not in self.preferencias_citas[vendedor]:
                                self.preferencias_citas[vendedor].append(comprador)
                
                # Procesar el resto del archivo
                for row in reader:
                    if len(row) >= 2:
                        vendedor = row[0].strip()
                        comprador = row[1].strip()
                        if vendedor and comprador:
                            vendedores_set.add(vendedor)
                            compradores_set.add(comprador)
                            if vendedor not in self.preferencias_citas:
                                self.preferencias_citas[vendedor] = []
                            if comprador not in self.preferencias_citas[vendedor]:
                                self.preferencias_citas[vendedor].append(comprador)
            
            # Actualizar listas de participantes con nombres reales
            self.vendedores = list(vendedores_set)
            self.compradores = list(compradores_set)
            
            # Completar con nombres genéricos si es necesario
            while len(self.vendedores) < self.num_vendedores:
                nuevo_nombre = f"Vendedor_Extra_{len(self.vendedores)+1:02d}"
                self.vendedores.append(nuevo_nombre)
            
            while len(self.compradores) < self.num_compradores:
                nuevo_nombre = f"Comprador_Extra_{len(self.compradores)+1:02d}"
                self.compradores.append(nuevo_nombre)
            
            # Inicializar contadores con los nombres reales
            self.citas_por_vendedor = {v: 0 for v in self.vendedores}
            self.citas_por_comprador = {c: 0 for c in self.compradores}
            
            # Inicializar tracking de encuentros para evitar repeticiones
            self.encuentros_realizados = {c: set() for c in self.compradores}
            
            # Actualizar números reales
            self.num_vendedores = len(self.vendedores)
            self.num_compradores = len(self.compradores)
            
            self.participantes_cargados = True
            
            total_preferencias = sum(len(compradores) for compradores in self.preferencias_citas.values())
            print(f"Cargadas {total_preferencias} preferencias de citas de {len(self.preferencias_citas)} vendedores desde {ruta_archivo}")
            print(f"Vendedores encontrados: {len(vendedores_set)}")
            print(f"Compradores encontrados: {len(compradores_set)}")
            return True
            
        except Exception as e:
            print(f"Error al cargar archivo de preferencias: {e}")
            return False

    def _inicializar_participantes_por_defecto(self):
        """Inicializa participantes con nombres genéricos si no se cargaron preferencias"""
        if not self.participantes_cargados:
            self.vendedores = [f"Vendedor_{i+1:02d}" for i in range(self.num_vendedores)]
            self.compradores = [f"Comprador_{i+1:02d}" for i in range(self.num_compradores)]
            self.citas_por_vendedor = {v: 0 for v in self.vendedores}
            self.citas_por_comprador = {c: 0 for c in self.compradores}
            self.participantes_cargados = True

    def _verificar_disponibilidad_horaria(self, participante: str, slot: int, es_vendedor: bool = True) -> bool:
        """Verifica si un participante está disponible en el horario del slot específico"""
        # NUEVA RESTRICCIÓN: Regional SAS solo disponible de 8:30 a 10:30 (slots 0-7 con 15 min cada uno desde 8:30)
        if participante == "REGIONAL S.A.S" and not es_vendedor:  # Es comprador
            return slot <= 7  # Slots 0-7 cubren desde 8:30 hasta 10:15
        
        # Restricciones horarias específicas reactivadas
        if participante == "NEIRA YORK COFFEE" and not es_vendedor:  # Es comprador
            # Ajustado: NEIRA YORK COFFEE de 10:15 a 12:00 (slots 7-13 con nuevo inicio 8:30)
            return 7 <= slot <= 13  # Slots 7-13 cubren desde 10:15 hasta 12:00
        
        if participante == "ENCADENAMIENTOS PRODUCTIVOS -  CAFE AROMAS DEL EJE / CAFÉ GRANEAO." and not es_vendedor:  # Es comprador
            # Ajustado: ENCADENAMIENTOS PRODUCTIVOS de 10:15 a 11:15 (slots 7-10 con nuevo inicio 8:30)
            return 7 <= slot <= 10  # Slots 7-10 cubren desde 10:15 hasta 11:15
        
        # NUEVA RESTRICCIÓN: Café Origen de la Montaña ES ENCADENAMIENTOS PRODUCTIVOS
        if participante == "CAFÉ ORIGEN DE LA MONTAÑA" and es_vendedor:  # Es vendedor
            # Debe actuar primero como comprador (ENCADENAMIENTOS), luego como vendedor
            # Como vendedor, disponible después de sus citas como comprador (después slot 10)
            return slot >= 11  # Disponible después de 11:15
        
        if participante == "La vuelta" and es_vendedor:  # Es vendedor
            # Ajustado: La vuelta de 12:00 a 13:00 (slots 14-17 con nuevo inicio 8:30)
            return slot >= 14  # Slots 14-17 cubren desde 12:00 hasta 13:00
        
        if participante == "Café Del Tajo" and es_vendedor:  # Es vendedor
            # Ajustado: Café Del Tajo de 11:15 a 13:00 (slots 11-17 con nuevo inicio 8:30)
            return slot >= 11  # Slots 11-17 cubren desde 11:15 hasta 13:00
        
        if participante == "Café Tradición Premium" and es_vendedor:  # Es vendedor
            # Ajustado: Café Tradición Premium de 11:15 a 13:00 (slots 11-17 con nuevo inicio 8:30)
            return slot >= 11  # Slots 11-17 cubren desde 11:15 hasta 13:00
        
        # Para todos los demás participantes, siempre disponibles
        return True

    def _puede_agendar_cita_grupo(self, vendedores: List[str], comprador: str, slot: int) -> bool:
        """Verifica si se puede agendar una cita grupal (3 vendedores + 1 comprador)"""
        # NUEVA VERIFICACIÓN: Restricciones horarias específicas
        # Verificar disponibilidad del comprador
        if not self._verificar_disponibilidad_horaria(comprador, slot, es_vendedor=False):
            return False
        
        # Verificar disponibilidad de cada vendedor
        for vendedor in vendedores:
            if not self._verificar_disponibilidad_horaria(vendedor, slot, es_vendedor=True):
                return False
        
        # RESTRICCIÓN ESPECÍFICA: ENCADENAMIENTOS PRODUCTIVOS no se reúne con Café Del Tajo ni Café Tradición Premium
        if comprador == "ENCADENAMIENTOS PRODUCTIVOS -  CAFE AROMAS DEL EJE / CAFÉ GRANEAO.":
            for vendedor in vendedores:
                if vendedor in ["Café Del Tajo", "Café Tradición Premium"]:
                    return False  # No permitir esta combinación
        
        # NUEVA RESTRICCIÓN: Verificar que solo hay 3 vendedores máximo por cita
        if len(vendedores) > 3:
            return False  # Máximo 3 vendedores por cita
        
        # Verificar límite de citas por vendedor
        for vendedor in vendedores:
            if vendedor not in self.vendedores:
                return False
            if self.citas_por_vendedor[vendedor] >= self.max_citas_vendedor:
                return False
        
        # NUEVA VERIFICACIÓN: No permitir que un vendedor tenga citas simultáneas con diferentes compradores
        for vendedor in vendedores:
            for cita_comprador, cita_vendedores in self.agenda[slot]:
                if vendedor in cita_vendedores and cita_comprador != comprador:
                    return False  # Este vendedor ya tiene una cita con otro comprador en este slot
        
        # Verificar que el comprador no esté ocupado en este slot
        for cita_comprador, cita_vendedores in self.agenda[slot]:
            if cita_comprador == comprador:
                return False  # Este comprador ya tiene una cita en este slot
        
        # NUEVA VERIFICACIÓN: No repetir encuentros comprador-vendedor
        if comprador in self.encuentros_realizados:
            for vendedor in vendedores:
                if vendedor in self.encuentros_realizados[comprador]:
                    return False  # Este comprador ya se reunió con este vendedor
        
        return True

    def _encontrar_vendedores_disponibles(self, slot: int, excluir: Set[str] = None) -> List[str]:
        """Encuentra vendedores disponibles para un slot específico"""
        if excluir is None:
            excluir = set()
        
        disponibles = []
        
        # Vendedores ya ocupados en este slot (cualquier cita en este slot los hace no disponibles)
        ocupados = set()
        for _, vendedores_en_slot in self.agenda[slot]:
            ocupados.update(vendedores_en_slot)
        
        for vendedor in self.vendedores:
            if (vendedor not in ocupados and 
                vendedor not in excluir and
                self.citas_por_vendedor[vendedor] < self.max_citas_vendedor and
                self._verificar_disponibilidad_horaria(vendedor, slot, es_vendedor=True)):  # NUEVA VERIFICACIÓN
                disponibles.append(vendedor)
        
        return disponibles

    def generar_agenda_optimizada(self) -> Dict:
        """Genera la agenda optimizada distribuyendo las citas equitativamente"""
        print("Generando agenda optimizada...")
        
        # Asegurarse de que los participantes estén inicializados
        self._inicializar_participantes_por_defecto()
        
        # Solo asignar citas con preferencias específicas (sin citas aleatorias)
        print("Procesando preferencias específicas...")
        self._asignar_citas_con_preferencias()
        
        # Validar que no hay conflictos de vendedores con citas simultáneas
        print("\nValidando agenda para conflictos de vendedores...")
        validacion_conflictos = self._validar_agenda_sin_conflictos()
        
        if validacion_conflictos["tiene_conflictos"]:
            print(f"⚠️  CONFLICTOS DETECTADOS: {validacion_conflictos['total_conflictos']} conflictos encontrados")
            print(f"   Vendedores afectados: {', '.join(validacion_conflictos['vendedores_afectados'])}")
            for conflicto in validacion_conflictos["conflictos"]:
                print(f"   🔴 {conflicto['horario']}: {', '.join(conflicto['vendedores_en_conflicto'])} en conflicto")
                print(f"      - Cita 1: {conflicto['cita1']['comprador']} con {conflicto['cita1']['vendedores']}")
                print(f"      - Cita 2: {conflicto['cita2']['comprador']} con {conflicto['cita2']['vendedores']}")
        else:
            print("✅ Validación exitosa: No se detectaron conflictos de vendedores con citas simultáneas")
        
        resultado = self._formatear_resultado()
        resultado["validacion_conflictos"] = validacion_conflictos
        
        return resultado

    def _asignar_citas_con_preferencias(self):
        """Asigna citas basadas en las preferencias cargadas"""
        # FASE 1: Concentrar TODOS los participantes en las primeras horas
        print("Fase 1: Concentrando TODOS los vendedores y compradores en las primeras horas...")
        self._garantizar_citas_iniciales_compradores()
        
        # FASE 2: Solo completar preferencias restantes si hay participantes sin asignar
        print("Fase 2: Completando citas con preferencias restantes (si es necesario)...")
        
        # Crear un mapa de todas las combinaciones vendedor-comprador preferidas
        combinaciones_preferidas = []
        for vendedor, compradores_pref in self.preferencias_citas.items():
            for comprador in compradores_pref:
                # Solo agregar si no se asignó ya en la fase 1
                if (vendedor, comprador) not in self.citas_preferencia_asignadas:
                    combinaciones_preferidas.append((vendedor, comprador))
        
        if not combinaciones_preferidas:
            print("✅ No hay preferencias adicionales por procesar - concentración completada")
            return
        
        # Agrupar por comprador para formar citas grupales
        vendedores_por_comprador = {}
        for vendedor, comprador in combinaciones_preferidas:
            if comprador not in vendedores_por_comprador:
                vendedores_por_comprador[comprador] = []
            vendedores_por_comprador[comprador].append(vendedor)
        
        # Asignar citas por comprador
        for comprador, vendedores_interesados in vendedores_por_comprador.items():
            # Mezclar para variedad
            random.shuffle(vendedores_interesados)
            
            # Dividir vendedores en grupos estratégicos: priorizar 5 citas con 2 vendedores + 5 citas con 3 vendedores
            while len(vendedores_interesados) >= 2:
                # Contar cuántas citas ya tiene este comprador para decidir estrategia
                citas_actuales_comprador = self.citas_por_comprador.get(comprador, 0)
                
                # Estrategia: 4 citas con 2 vendedores, luego citas con 3 vendedores
                if citas_actuales_comprador < 4 and len(vendedores_interesados) >= 2:
                    # Primeras 4 citas: usar 2 vendedores
                    grupo_vendedores = vendedores_interesados[:2]
                    vendedores_interesados = vendedores_interesados[2:]
                elif len(vendedores_interesados) >= 3:
                    # Siguientes citas: usar 3 vendedores
                    grupo_vendedores = vendedores_interesados[:3]
                    vendedores_interesados = vendedores_interesados[3:]
                else:
                    # Usar lo que quede disponible
                    grupo_vendedores = vendedores_interesados[:2]
                    vendedores_interesados = vendedores_interesados[2:]
                
                # Buscar slot disponible para este grupo
                slot_asignado = self._buscar_slot_disponible(grupo_vendedores, comprador)
                if slot_asignado is not None:
                    self.agenda[slot_asignado].append((comprador, grupo_vendedores))
                    for vendedor in grupo_vendedores:
                        self.citas_por_vendedor[vendedor] += 1
                        self.citas_preferencia_asignadas.add((vendedor, comprador))
                        # Registrar encuentro realizado
                        self.encuentros_realizados[comprador].add(vendedor)
                    self.citas_por_comprador[comprador] += 1
                    print(f"✓ Cita preferida asignada: {grupo_vendedores} → {comprador} ({len(grupo_vendedores)} vendedores)")
            
            # Si quedan vendedores (1 o pocos), intentar completar grupos existentes
            if vendedores_interesados:
                vendedores_adicionales = self._encontrar_vendedores_para_completar_grupo(
                    vendedores_interesados, comprador
                )
                if vendedores_adicionales:
                    grupo_completo = vendedores_interesados + vendedores_adicionales
                    slot_asignado = self._buscar_slot_disponible(grupo_completo, comprador)
                    if slot_asignado is not None:
                        self.agenda[slot_asignado].append((comprador, grupo_completo))
                        for vendedor in grupo_completo:
                            self.citas_por_vendedor[vendedor] += 1
                            if vendedor in vendedores_interesados:
                                self.citas_preferencia_asignadas.add((vendedor, comprador))
                            # Registrar encuentro realizado
                            self.encuentros_realizados[comprador].add(vendedor)
                        self.citas_por_comprador[comprador] += 1
                        print(f"✓ Cita preferida completada: {grupo_completo} → {comprador}")
                elif len(vendedores_interesados) == 1:
                    # Si solo queda 1 vendedor, crear una cita de 1 vendedor
                    slot_asignado = self._buscar_slot_disponible(vendedores_interesados, comprador)
                    if slot_asignado is not None:
                        self.agenda[slot_asignado].append((comprador, vendedores_interesados))
                        for vendedor in vendedores_interesados:
                            self.citas_por_vendedor[vendedor] += 1
                            self.citas_preferencia_asignadas.add((vendedor, comprador))
                            self.encuentros_realizados[comprador].add(vendedor)
                        self.citas_por_comprador[comprador] += 1
                        print(f"✓ Cita preferida individual: {vendedores_interesados} → {comprador}")
                else:
                    # Si no se pueden completar, intentar asignar estos vendedores a otras citas
                    self._asignar_vendedores_restantes(vendedores_interesados, comprador)

    def _procesar_citas_criticas(self):
        """Procesa citas críticas que DEBEN darse en las primeras horas"""
        print("🔥 PROCESANDO CITAS CRÍTICAS EN PRIMERAS HORAS...")
        
        # Citas críticas que DEBEN darse en slots 0-3 (08:30-09:15)
        citas_criticas = [
            ("D'CLEO COFFEE", "BOX BRAND", 0),  # Slot 0: 08:30-08:45
        ]
        
        for vendedor, comprador, slot_preferido in citas_criticas:
            if vendedor in self.vendedores and comprador in self.compradores:
                # Verificar disponibilidad en el slot preferido
                if (self._verificar_disponibilidad_horaria(vendedor, slot_preferido, es_vendedor=True) and
                    self._verificar_disponibilidad_horaria(comprador, slot_preferido, es_vendedor=False)):
                    
                    # Buscar si ya existe una cita para este comprador en este slot
                    cita_existente = None
                    for i, (comp_existente, vendedores_existentes) in enumerate(self.agenda[slot_preferido]):
                        if comp_existente == comprador:
                            cita_existente = i
                            break
                    
                    if cita_existente is not None:
                        # Agregar vendedor a cita existente si hay espacio
                        vendedores_actuales = self.agenda[slot_preferido][cita_existente][1]
                        if len(vendedores_actuales) < 3 and vendedor not in vendedores_actuales:
                            vendedores_actuales.append(vendedor)
                            # Actualizar contadores
                            self.citas_por_vendedor[vendedor] += 1
                            self.citas_preferencia_asignadas.add((vendedor, comprador))
                            if comprador not in self.encuentros_realizados:
                                self.encuentros_realizados[comprador] = set()
                            self.encuentros_realizados[comprador].add(vendedor)
                            print(f"   ✅ CRÍTICA: {vendedor} agregado a cita existente con {comprador} en {self.horarios[slot_preferido]}")
                        else:
                            print(f"   ⚠️  CRÍTICA: No se pudo agregar {vendedor} a {comprador} (cita llena)")
                    else:
                        # Crear nueva cita para el comprador crítico
                        nueva_cita = (comprador, [vendedor])
                        self.agenda[slot_preferido].append(nueva_cita)
                        # Actualizar contadores
                        self.citas_por_vendedor[vendedor] += 1
                        self.citas_por_comprador[comprador] += 1
                        self.citas_preferencia_asignadas.add((vendedor, comprador))
                        if comprador not in self.encuentros_realizados:
                            self.encuentros_realizados[comprador] = set()
                        self.encuentros_realizados[comprador].add(vendedor)
                        print(f"   ✅ CRÍTICA: Nueva cita {vendedor} ↔ {comprador} en {self.horarios[slot_preferido]}")
                else:
                    print(f"   ❌ CRÍTICA: {vendedor} ↔ {comprador} no disponible en {self.horarios[slot_preferido]}")

    def _garantizar_citas_iniciales_compradores(self):
        """Garantiza distribución estratégica con horario extendido y preferencia por 3 vendedores por cita"""
        print("🚀 DISTRIBUCIÓN ESTRATÉGICA ACTUALIZADA:")
        print("🎯 08:30-10:30: Regional SAS disponible - máximo 3 vendedores por cita")
        print("🎯 10:15-11:15: ENCADENAMIENTOS PRODUCTIVOS como comprador - máximo 3 vendedores por cita")
        print("🎯 11:15-13:00: Café Origen de la Montaña como vendedor - máximo 3 vendedores por cita")
        print("⏰ Restricciones horarias:")
        print("   • REGIONAL S.A.S (08:30-10:30)")
        print("   • NEIRA YORK COFFEE (10:15-12:00)")
        print("   • ENCADENAMIENTOS PRODUCTIVOS (10:15-11:15)")
        print("   • Café Origen de la Montaña como vendedor (11:15-13:00)")
        print("   • La vuelta (12:00-13:00)")
        print("   • Café Del Tajo (11:15-13:00)")
        print("   • Café Tradición Premium (11:15-13:00)")
        print("🚫 Restricciones específicas:")
        print("   • ENCADENAMIENTOS PRODUCTIVOS no se reúne con Café Del Tajo ni Café Tradición Premium")
        print("   • Ningún vendedor puede tener citas simultáneas con diferentes compradores")
        print("   • Máximo 3 vendedores por cita por comprador")
        print("   • Café Origen de la Montaña: primero como comprador (ENCADENAMIENTOS), luego como vendedor")
        print("   • 🔥 CRÍTICA: D'CLEO COFFEE ↔ BOX BRAND debe darse en las primeras horas (08:30-09:15)")
        
        # PRIMERA PRIORIDAD: Asegurar citas críticas en los primeros slots
        self._procesar_citas_criticas()
        
        # Usar SOLO los primeros 2 slots para concentrar a TODOS los compradores
        slots_iniciales = 2
        compradores_sin_cita = set(self.compradores)
        vendedores_sin_cita = set(self.vendedores)
        
        print(f"📅 Usando SOLO slots 1-{slots_iniciales} para concentrar TODOS los compradores...")
        
        # FASE 1: Distribución estratégica por slot
        for slot in range(slots_iniciales):
            print(f"\n⏰ PROCESANDO SLOT {slot+1} ({self.horarios[slot]})...")
            
            # Determinar estrategia según el horario
            # Slots 0-3: 10:15-11:15 (4 citas 2v + 6 citas 3v = 26 vendedores)
            # Slots 4+: 11:15-13:00 (4 citas 2v + 5 citas 3v = 23 vendedores)
            if slot <= 3:  # 10:15-11:15
                target_citas_2v = 4  # Meta: 4 citas con 2 vendedores
                target_citas_3v = 6  # Meta: 6 citas con 3 vendedores
                total_vendedores_meta = 26
                periodo = "10:15-11:15"
            else:  # 11:15-13:00
                target_citas_2v = 4  # Meta: 4 citas con 2 vendedores
                target_citas_3v = 5  # Meta: 5 citas con 3 vendedores
                total_vendedores_meta = 23
                periodo = "11:15-13:00"
            
            print(f"   🎯 Meta ({periodo}): {target_citas_2v} citas (2v) + {target_citas_3v} citas (3v) = {total_vendedores_meta} vendedores")
            
            # Lista de compradores disponibles para este slot (considerando restricciones horarias)
            compradores_disponibles = [c for c in compradores_sin_cita 
                                     if not any(comp == c for comp, _ in self.agenda[slot])
                                     and self._verificar_disponibilidad_horaria(c, slot, es_vendedor=False)]
            
            # Planificar distribución según el período
            citas_planificadas = []
            citas_con_2_vendedores = 0
            citas_con_3_vendedores = 0
            
            # Asignar compradores con distribución estratégica
            for i, comprador in enumerate(compradores_disponibles[:10]):  # Máximo 10 compradores
                if comprador in compradores_sin_cita:
                    # Buscar vendedores con preferencias para este comprador
                    vendedores_preferidos = []
                    for vendedor, compradores_pref in self.preferencias_citas.items():
                        if (comprador in compradores_pref and vendedor in vendedores_sin_cita and
                            self._verificar_disponibilidad_horaria(vendedor, slot, es_vendedor=True)):
                            vendedores_preferidos.append(vendedor)
                    
                    # Obtener vendedores disponibles para este slot
                    vendedores_disponibles_slot = self._encontrar_vendedores_disponibles(slot)
                    
                    # Combinar preferidos + disponibles (priorizando preferidos)
                    candidatos = []
                    # Primero los preferidos que están disponibles
                    for v in vendedores_preferidos:
                        if v in vendedores_disponibles_slot and v not in self.encuentros_realizados[comprador]:
                            candidatos.append(v)
                    
                    # Luego otros vendedores disponibles
                    for v in vendedores_disponibles_slot:
                        if (v not in candidatos and 
                            v not in self.encuentros_realizados[comprador] and
                            v in vendedores_sin_cita):
                            candidatos.append(v)
                    
                    # Si no hay suficientes sin repetir, usar cualquier disponible
                    if len(candidatos) < 2:
                        for v in vendedores_disponibles_slot:
                            if v not in candidatos and v in vendedores_sin_cita:
                                candidatos.append(v)
                    
                    # ESTRATEGIA DE DISTRIBUCIÓN: Alternar entre 2 y 3 vendedores
                    if candidatos:
                        # Decidir cuántos vendedores asignar basado en la meta estratégica
                        if citas_con_2_vendedores < target_citas_2v and len(candidatos) >= 2:
                            # Asignar 2 vendedores (primeras 5 citas)
                            num_vendedores = 2
                            tipo_cita = "2v"
                        elif citas_con_3_vendedores < target_citas_3v and len(candidatos) >= 3:
                            # Asignar 3 vendedores (siguientes 5 citas)
                            num_vendedores = 3
                            tipo_cita = "3v"
                        elif len(candidatos) >= 2:
                            # Si no podemos cumplir la meta exacta, usar lo que esté disponible
                            if citas_con_2_vendedores < citas_con_3_vendedores:
                                num_vendedores = 2
                                tipo_cita = "2v"
                            else:
                                num_vendedores = min(3, len(candidatos))
                                tipo_cita = "3v"
                        else:
                            # Emergencia: usar lo que haya disponible
                            num_vendedores = len(candidatos)
                            tipo_cita = f"{num_vendedores}v"
                        
                        grupo_vendedores = candidatos[:num_vendedores]
                        
                        # Verificar que se puede agendar
                        if self._puede_agendar_cita_grupo(grupo_vendedores, comprador, slot):
                            self.agenda[slot].append((comprador, grupo_vendedores))
                            
                            # Actualizar contadores
                            for vendedor in grupo_vendedores:
                                self.citas_por_vendedor[vendedor] += 1
                                self.citas_preferencia_asignadas.add((vendedor, comprador))
                                self.encuentros_realizados[comprador].add(vendedor)
                                vendedores_sin_cita.discard(vendedor)
                            
                            self.citas_por_comprador[comprador] += 1
                            compradores_sin_cita.discard(comprador)
                            
                            # Actualizar contadores de distribución
                            if num_vendedores == 2:
                                citas_con_2_vendedores += 1
                            elif num_vendedores == 3:
                                citas_con_3_vendedores += 1
                            
                            vendedores_usados = len(grupo_vendedores)
                            print(f"   ✅ {comprador} ↔ [{', '.join(grupo_vendedores)}] ({num_vendedores} vendedores - {tipo_cita})")
            
            # Reporte del slot
            total_vendedores_usados = sum(len(vendedores) for _, vendedores in self.agenda[slot])
            print(f"   📊 Slot {slot+1} completado:")
            print(f"      • Citas con 2 vendedores: {citas_con_2_vendedores}/{target_citas_2v}")
            print(f"      • Citas con 3 vendedores: {citas_con_3_vendedores}/{target_citas_3v}")
            print(f"      • Total vendedores usados: {total_vendedores_usados}/{total_vendedores_meta}")
            print(f"      • Total citas: {len(self.agenda[slot])}/10")
            
            # Si ya todos los compradores tienen cita, salir del bucle
            if not compradores_sin_cita:
                print(f"   🎉 ¡TODOS los compradores asignados en primeros {slot+1} slots!")
                break
        
        # EMERGENCIA: Si aún hay compradores sin cita después de 2 slots, forzar asignación
        if compradores_sin_cita:
            print(f"\n� MODO EMERGENCIA: {len(compradores_sin_cita)} compradores sin cita, forzando asignación...")
            
            for comprador in list(compradores_sin_cita):
                # Buscar cualquier slot de los primeros 2 donde podamos meter este comprador
                for slot in range(slots_iniciales):
                    if (comprador in compradores_sin_cita and 
                        self._verificar_disponibilidad_horaria(comprador, slot, es_vendedor=False)):  # Verificar restricción horaria
                        
                        # Buscar vendedores disponibles (incluso si solo hay 1)
                        vendedores_disponibles_slot = self._encontrar_vendedores_disponibles(slot)
                        
                        # Filtrar vendedores que no se hayan reunido con este comprador
                        candidatos = []
                        for v in vendedores_disponibles_slot:
                            if v not in self.encuentros_realizados[comprador]:
                                candidatos.append(v)
                        
                        # Si no hay sin repetir, usar cualquier disponible
                        if not candidatos:
                            candidatos = vendedores_disponibles_slot
                        
                        if candidatos:
                            # Usar estrategia balanceada incluso en emergencia según el período
                            citas_existentes = len(self.agenda[slot])
                            
                            # Aplicar distribución según el período
                            if slot <= 3:  # 10:15-11:15: 4 citas 2v + 6 citas 3v
                                if citas_existentes < 4:
                                    # Primeras 4 citas: preferir 2 vendedores
                                    num_vendedores = min(2, len(candidatos))
                                else:
                                    # Siguientes 6 citas: permitir 3 vendedores
                                    num_vendedores = min(3, len(candidatos))
                            else:  # 11:15-13:00: 4 citas 2v + 5 citas 3v
                                if citas_existentes < 4:
                                    # Primeras 4 citas: preferir 2 vendedores
                                    num_vendedores = min(2, len(candidatos))
                                else:
                                    # Siguientes 5 citas: permitir 3 vendedores
                                    num_vendedores = min(3, len(candidatos))
                            
                            grupo_vendedores = candidatos[:num_vendedores]
                            
                            # Verificar disponibilidad final
                            if self._puede_agendar_cita_grupo(grupo_vendedores, comprador, slot):
                                self.agenda[slot].append((comprador, grupo_vendedores))
                                
                                # Actualizar contadores
                                for vendedor in grupo_vendedores:
                                    self.citas_por_vendedor[vendedor] += 1
                                    self.citas_preferencia_asignadas.add((vendedor, comprador))
                                    self.encuentros_realizados[comprador].add(vendedor)
                                    vendedores_sin_cita.discard(vendedor)
                                
                                self.citas_por_comprador[comprador] += 1
                                compradores_sin_cita.discard(comprador)
                                
                                print(f"   🆘 EMERGENCIA: {comprador} ↔ [{', '.join(grupo_vendedores)}] (slot {slot+1})")
                                break  # Salir del bucle de slots para este comprador
        
        # Reportes finales
        compradores_con_cita = len(self.compradores) - len(compradores_sin_cita)
        vendedores_con_cita = len(self.vendedores) - len(vendedores_sin_cita)
        
        print(f"\n📊 REPORTE FINAL DE DISTRIBUCIÓN ESTRATÉGICA:")
        print(f"   👥 Compradores asignados: {compradores_con_cita}/{len(self.compradores)}")
        print(f"   🏪 Vendedores asignados: {vendedores_con_cita}/{len(self.vendedores)}")
        
        # Calcular distribución real por slot
        for slot in range(slots_iniciales):
            if self.agenda[slot]:
                citas_2v = sum(1 for _, vendedores in self.agenda[slot] if len(vendedores) == 2)
                citas_3v = sum(1 for _, vendedores in self.agenda[slot] if len(vendedores) == 3)
                total_vendedores_slot = sum(len(vendedores) for _, vendedores in self.agenda[slot])
                
                # Mostrar metas según el período
                if slot <= 3:  # 10:15-11:15
                    meta_descripcion = "4 citas (2v) + 6 citas (3v) = 26 vendedores"
                else:  # 11:15-13:00
                    meta_descripcion = "4 citas (2v) + 5 citas (3v) = 23 vendedores"
                
                print(f"   📈 Slot {slot+1}: {citas_2v} citas (2v) + {citas_3v} citas (3v) = {total_vendedores_slot} vendedores (Meta: {meta_descripcion})")
        
        if compradores_sin_cita:
            print(f"   ⚠️  Compradores pendientes: {list(compradores_sin_cita)}")
        
        if vendedores_sin_cita:
            print(f"   ⚠️  Vendedores pendientes: {list(vendedores_sin_cita)}")
            
        if not compradores_sin_cita:
            print(f"   🎉 ¡ÉXITO! TODOS los compradores asignados con distribución estratégica!")
        
        # Mostrar distribución por slot
        print(f"\n📈 DISTRIBUCIÓN FINAL EN {slots_iniciales} SLOTS:")
        for slot in range(slots_iniciales):
            num_citas = len(self.agenda[slot])
            if num_citas > 0:
                print(f"   Slot {slot+1} ({self.horarios[slot]}): {num_citas} citas")
            else:
                print(f"   Slot {slot+1} ({self.horarios[slot]}): Sin citas")

    def _buscar_slot_disponible(self, vendedores: List[str], comprador: str) -> Optional[int]:
        """Busca un slot disponible para el grupo de vendedores y comprador"""
        for slot in range(self.num_slots):
            if self._puede_agendar_cita_grupo(vendedores, comprador, slot):
                return slot
        return None

    def _asignar_vendedores_restantes(self, vendedores_restantes: List[str], comprador_preferido: str):
        """Intenta asignar vendedores restantes a citas existentes o nuevas"""
        for vendedor in vendedores_restantes:
            asignado = False
            
            # Intentar agregar a una cita existente del comprador preferido que tenga espacio
            for slot in range(self.num_slots):
                for i, (comprador, vendedores_en_cita) in enumerate(self.agenda[slot]):
                    if (comprador == comprador_preferido and 
                        len(vendedores_en_cita) < 3 and  # Máximo 3, pero preferir 2
                        vendedor not in vendedores_en_cita and
                        vendedor not in self.encuentros_realizados[comprador] and  # Nueva verificación
                        self.citas_por_vendedor[vendedor] < self.max_citas_vendedor):
                        
                        # Solo agregar si la cita tiene 1 vendedor (para llegar a 2) o si hay mucha demanda
                        if len(vendedores_en_cita) == 1 or (len(vendedores_en_cita) == 2 and self.citas_por_vendedor[vendedor] < 2):
                            
                            # Verificar que el vendedor no esté ocupado en este slot
                            vendedor_ocupado_en_slot = any(
                                vendedor in otros_vendedores 
                                for _, otros_vendedores in self.agenda[slot]
                            )
                            
                            if not vendedor_ocupado_en_slot:
                                self.agenda[slot][i] = (comprador, vendedores_en_cita + [vendedor])
                                self.citas_por_vendedor[vendedor] += 1
                                self.citas_preferencia_asignadas.add((vendedor, comprador))
                                # Registrar encuentro realizado
                                self.encuentros_realizados[comprador].add(vendedor)
                                print(f"✓ Vendedor agregado a cita existente: {vendedor} → {comprador}")
                                asignado = True
                                break
                if asignado:
                    break

    def _encontrar_vendedores_para_completar_grupo(self, vendedores_base: List[str], comprador: str) -> List[str]:
        """Encuentra vendedores adicionales para completar un grupo de 2 (o 3 si es necesario)"""
        # Priorizar grupos de 2, completar a 3 solo si hay muchos vendedores disponibles
        if len(vendedores_base) >= 2:
            return []  # Ya tenemos suficientes con 2
        
        necesarios = 2 - len(vendedores_base)  # Completar hasta 2 por defecto
        candidatos = []
        
        for vendedor in self.vendedores:
            if (vendedor not in vendedores_base and
                vendedor not in self.encuentros_realizados[comprador] and  # Nueva verificación
                self.citas_por_vendedor[vendedor] < self.max_citas_vendedor):
                candidatos.append(vendedor)
        
        # Si hay muchos vendedores disponibles (más de 4), podemos permitir grupos de 3
        if len(candidatos) > 4 and len(vendedores_base) == 2:
            necesarios = 1  # Agregar uno más para hacer grupo de 3
        
        return candidatos[:necesarios]

    def _validar_agenda_sin_conflictos(self) -> Dict[str, any]:
        """Valida que la agenda final no tenga conflictos de vendedores con citas simultáneas"""
        conflictos_detectados = []
        vendedores_con_conflictos = set()
        
        # Revisar cada slot para detectar conflictos
        for slot in range(self.num_slots):
            citas_en_slot = self.agenda[slot]
            
            # Verificar cada par de citas en el mismo slot
            for i, (comprador1, vendedores1) in enumerate(citas_en_slot):
                for j, (comprador2, vendedores2) in enumerate(citas_en_slot):
                    if i < j:  # Evitar comparar la misma cita dos veces
                        # Buscar vendedores que aparecen en ambas citas
                        vendedores_repetidos = set(vendedores1) & set(vendedores2)
                        if vendedores_repetidos:
                            conflicto = {
                                "slot": slot,
                                "horario": self.horarios[slot],
                                "vendedores_en_conflicto": list(vendedores_repetidos),
                                "cita1": {"comprador": comprador1, "vendedores": vendedores1},
                                "cita2": {"comprador": comprador2, "vendedores": vendedores2}
                            }
                            conflictos_detectados.append(conflicto)
                            vendedores_con_conflictos.update(vendedores_repetidos)
        
        return {
            "tiene_conflictos": len(conflictos_detectados) > 0,
            "total_conflictos": len(conflictos_detectados),
            "conflictos": conflictos_detectados,
            "vendedores_afectados": list(vendedores_con_conflictos)
        }

    def _completar_agenda_restante(self):
        """Completa la agenda manteniendo la estrategia: 4 citas (2v) + 6 citas (3v) antes 11:15, 4 citas (2v) + 5 citas (3v) después"""
        for slot in range(self.num_slots):
            # Para cada slot, intentar agendar citas manteniendo distribución estratégica
            citas_actuales_slot = len(self.agenda[slot])
            
            # Determinar metas según el período
            if slot <= 3:  # 10:15-11:15: 4 citas 2v + 6 citas 3v
                target_citas_2v = 4
                target_citas_3v = 6
            else:  # 11:15-13:00: 4 citas 2v + 5 citas 3v
                target_citas_2v = 4
                target_citas_3v = 5
            
            for comprador in self.compradores:
                # Verificar si este comprador ya tiene una cita en este slot
                comprador_ocupado = any(comp == comprador for comp, _ in self.agenda[slot])
                
                if not comprador_ocupado and citas_actuales_slot < 10:  # Máximo 10 citas por slot
                    # Encontrar vendedores disponibles para este slot
                    vendedores_disponibles = self._encontrar_vendedores_disponibles(slot)
                    
                    # Filtrar vendedores que no se hayan reunido con este comprador
                    vendedores_no_repetidos = [
                        v for v in vendedores_disponibles 
                        if v not in self.encuentros_realizados[comprador]
                    ]
                    
                    # Decidir número de vendedores según estrategia de distribución
                    if len(vendedores_no_repetidos) >= 2:
                        # Contar citas existentes en este slot por tipo
                        citas_2v = sum(1 for _, vendedores in self.agenda[slot] if len(vendedores) == 2)
                        citas_3v = sum(1 for _, vendedores in self.agenda[slot] if len(vendedores) == 3)
                        
                        # Estrategia según el período
                        if citas_2v < target_citas_2v and len(vendedores_no_repetidos) >= 2:
                            # Priorizar citas con 2 vendedores (primeras 4)
                            num_vendedores = 2
                        elif citas_3v < target_citas_3v and len(vendedores_no_repetidos) >= 3:
                            # Luego citas con 3 vendedores (6 o 5 según período)
                            num_vendedores = 3
                        elif len(vendedores_no_repetidos) >= 2:
                            # Si ya se cumplió la meta, usar lo que sea más eficiente
                            num_vendedores = 2  # Preferir 2 por eficiencia
                        else:
                            num_vendedores = 1  # Solo si no hay más opciones
                        
                        grupo_vendedores = vendedores_no_repetidos[:num_vendedores]
                        
                        # Verificar que se puede agendar
                        if self._puede_agendar_cita_grupo(grupo_vendedores, comprador, slot):
                            self.agenda[slot].append((comprador, grupo_vendedores))
                            for vendedor in grupo_vendedores:
                                self.citas_por_vendedor[vendedor] += 1
                                # Registrar encuentro realizado
                                self.encuentros_realizados[comprador].add(vendedor)
                            self.citas_por_comprador[comprador] += 1
                            citas_actuales_slot += 1

    def _encontrar_comprador_disponible(self, slot: int) -> Optional[str]:
        """Encuentra un comprador disponible para el slot"""
        # Compradores ya ocupados en este slot
        ocupados = {comprador for comprador, _ in self.agenda[slot]}
        
        # Buscar comprador con menos citas asignadas
        comprador_menos_citas = None
        min_citas = float('inf')
        
        for comprador in self.compradores:
            if comprador not in ocupados and self.citas_por_comprador[comprador] < min_citas:
                min_citas = self.citas_por_comprador[comprador]
                comprador_menos_citas = comprador
        
        return comprador_menos_citas

    def _formatear_resultado(self) -> Dict:
        """Formatea el resultado de la agenda"""
        resultado = {
            "configuracion": {
                "vendedores": self.num_vendedores,
                "compradores": self.num_compradores,
                "duracion_total": f"{(self.fin - self.inicio).total_seconds() / 60:.0f} minutos",
                "duracion_cita": f"{self.duracion_cita} minutos",
                "slots_disponibles": self.num_slots,
                "max_citas_por_vendedor": self.max_citas_vendedor,
                "vendedores_por_cita": self.vendedores_por_cita,
                "horario_inicio": self.inicio.strftime("%H:%M"),
                "horario_fin": self.fin.strftime("%H:%M"),
                "preferencias_cargadas": sum(len(compradores) for compradores in self.preferencias_citas.values())
            },
            "agenda": {},
            "estadisticas": {},
            "resumen_por_vendedor": {},
            "resumen_por_comprador": {},
            "preferencias_cumplidas": {}
        }
        
        # Agenda por horarios
        total_citas = 0
        total_encuentros = 0  # Total de encuentros individuales (vendedor-comprador)
        for slot, citas in self.agenda.items():
            if citas:
                citas_formateadas = []
                for comprador, vendedores in citas:
                    citas_formateadas.append({
                        "comprador": comprador,
                        "vendedores": vendedores
                    })
                    total_encuentros += len(vendedores)  # 3 encuentros por cita
                resultado["agenda"][self.horarios[slot]] = citas_formateadas
                total_citas += len(citas)
        
        # Estadísticas generales
        resultado["estadisticas"] = {
            "total_citas_programadas": total_citas,
            "total_encuentros_individuales": total_encuentros,
            "porcentaje_utilizacion_slots": f"{(total_citas / self.num_slots) * 100:.1f}%",
            "citas_promedio_por_vendedor": f"{sum(self.citas_por_vendedor.values()) / self.num_vendedores:.1f}",
            "citas_promedio_por_comprador": f"{sum(self.citas_por_comprador.values()) / self.num_compradores:.1f}",
            "preferencias_cumplidas": len(self.citas_preferencia_asignadas),
            "total_preferencias": sum(len(compradores) for compradores in self.preferencias_citas.values())
        }
        
        # Resumen por vendedor
        for vendedor in self.vendedores:
            citas_vendedor = []
            for slot, citas in self.agenda.items():
                for comprador, vendedores in citas:
                    if vendedor in vendedores:
                        citas_vendedor.append({
                            "horario": self.horarios[slot],
                            "comprador": comprador,
                            "otros_vendedores": [v for v in vendedores if v != vendedor]
                        })
            
            resultado["resumen_por_vendedor"][vendedor] = {
                "total_citas": len(citas_vendedor),
                "citas": citas_vendedor
            }
        
        # Resumen por comprador
        for comprador in self.compradores:
            citas_comprador = []
            for slot, citas in self.agenda.items():
                for comp, vendedores in citas:
                    if comp == comprador:
                        citas_comprador.append({
                            "horario": self.horarios[slot],
                            "vendedores": vendedores
                        })
            
            resultado["resumen_por_comprador"][comprador] = {
                "total_citas": len(citas_comprador),
                "citas": citas_comprador
            }
        
        # Preferencias cumplidas
        preferencias_totales = sum(len(compradores) for compradores in self.preferencias_citas.values())
        for vendedor, compradores_pref in self.preferencias_citas.items():
            vendedor_preferencias = []
            for comprador in compradores_pref:
                cumplida = (vendedor, comprador) in self.citas_preferencia_asignadas
                vendedor_preferencias.append({
                    "comprador": comprador,
                    "cumplida": cumplida
                })
            resultado["preferencias_cumplidas"][vendedor] = vendedor_preferencias
        
        return resultado

    def imprimir_agenda(self, resultado: Dict):
        """Imprime la agenda de forma legible"""
        print("\n" + "="*80)
        print("AGENDA DE RUEDA DE NEGOCIOS")
        print("="*80)
        
        print(f"\nESTADÍSTICAS:")
        stats = resultado["estadisticas"]
        print(f"• Total de citas programadas: {stats['total_citas_programadas']}")
        print(f"• Total encuentros individuales: {stats['total_encuentros_individuales']}")
        print(f"• Utilización de slots: {stats['porcentaje_utilizacion_slots']}")
        print(f"• Promedio citas por vendedor: {stats['citas_promedio_por_vendedor']}")
        print(f"• Promedio citas por comprador: {stats['citas_promedio_por_comprador']}")
        print(f"• Preferencias cumplidas: {stats['preferencias_cumplidas']}/{stats['total_preferencias']}")
        
        # Matriz Compradores-Horarios resumida
        print(f"\nMATRIZ COMPRADORES-HORARIOS:")
        print("-" * 80)
        
        # Crear matriz resumida para mostrar
        compradores_activos = [c for c in self.compradores if self.citas_por_comprador[c] > 0]
        horarios_con_citas = []
        
        for slot in range(self.num_slots):
            if len(self.agenda[slot]) > 0:
                horarios_con_citas.append(self.horarios[slot])
        
        # Mostrar solo primeros 6 horarios con citas para el resumen
        horarios_muestra = horarios_con_citas[:6] if len(horarios_con_citas) > 6 else horarios_con_citas
        
        # Encabezado de la matriz
        print(f"{'COMPRADOR':<25}", end="")
        for horario in horarios_muestra:
            print(f"{horario:>15}", end="")
        if len(horarios_con_citas) > 6:
            print(f"{'...':>15}", end="")
        print()
        print("-" * (25 + len(horarios_muestra) * 15 + (15 if len(horarios_con_citas) > 6 else 0)))
        
        # Filas de compradores
        for comprador in compradores_activos[:8]:  # Mostrar solo primeros 8 compradores
            print(f"{comprador[:24]:<25}", end="")
            
            for horario in horarios_muestra:
                # Buscar vendedores para este comprador en este horario
                vendedores_en_horario = []
                slot_idx = self.horarios.index(horario) if horario in self.horarios else -1
                
                if slot_idx >= 0 and slot_idx < len(self.agenda):
                    for cita_comprador, cita_vendedores in self.agenda[slot_idx]:
                        if cita_comprador == comprador:
                            vendedores_en_horario.extend(cita_vendedores)
                
                # Mostrar vendedores (máximo 2 para que quepa)
                if vendedores_en_horario:
                    vendedores_texto = ", ".join([v[:8] for v in vendedores_en_horario[:2]])
                    if len(vendedores_en_horario) > 2:
                        vendedores_texto += "..."
                    print(f"{vendedores_texto[:14]:>15}", end="")
                else:
                    print(f"{'':>15}", end="")
            
            if len(horarios_con_citas) > 6:
                print(f"{'':>15}", end="")
            print()
        
        if len(compradores_activos) > 8:
            print(f"{'... y ' + str(len(compradores_activos) - 8) + ' más':<25}")
        
        print(f"\n💡 Matriz completa disponible en: agenda_rueda_negocios.xlsx (Hoja: Matriz Compradores-Horarios)")
        
        print(f"\nAGENDA POR HORARIOS:")
        print("-"*80)
        
        for horario, citas in resultado["agenda"].items():
            print(f"\n{horario}")
            print("-" * 50)
            for i, cita in enumerate(citas, 1):
                vendedores_str = ", ".join(cita['vendedores'])
                print(f"  Cita {i}: {cita['comprador']} ↔ [{vendedores_str}]")
        
        # Mostrar preferencias cumplidas
        if resultado.get("preferencias_cumplidas"):
            print(f"\nPREFERENCIAS DE CITAS:")
            print("-"*50)
            cumplidas_total = 0
            total_preferencias = 0
            
            for vendedor, preferencias_vendedor in resultado["preferencias_cumplidas"].items():
                for pref in preferencias_vendedor:
                    total_preferencias += 1
                    estado = "✓" if pref["cumplida"] else "✗"
                    print(f"  {estado} {vendedor} → {pref['comprador']}")
                    if pref["cumplida"]:
                        cumplidas_total += 1
            
            print(f"\nTotal cumplidas: {cumplidas_total}/{total_preferencias}")

    def exportar_a_csv(self, resultado: Dict, nombre_archivo: str = "agenda_rueda_negocios.csv"):
        """Exporta la agenda a un archivo CSV"""
        ruta_archivo = f"c:\\Users\\angel\\OneDrive\\Escritorio\\escritorio\\Angela\\Programación agenda\\{nombre_archivo}"
        
        with open(ruta_archivo, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            
            # Encabezados
            writer.writerow(['Horario', 'Comprador', 'Vendedor_1', 'Vendedor_2', 'Vendedor_3'])
            
            # Datos
            for horario, citas in resultado["agenda"].items():
                for cita in citas:
                    vendedores = cita['vendedores'] + [''] * (3 - len(cita['vendedores']))  # Asegurar 3 columnas
                    writer.writerow([horario, cita['comprador']] + vendedores[:3])
        
        print(f"\nAgenda exportada a: {ruta_archivo}")

    def exportar_resumen_vendedores(self, resultado: Dict, nombre_archivo: str = "resumen_vendedores.csv"):
        """Exporta el resumen por vendedores a CSV"""
        ruta_archivo = f"c:\\Users\\angel\\OneDrive\\Escritorio\\escritorio\\Angela\\Programación agenda\\{nombre_archivo}"
        
        with open(ruta_archivo, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            
            # Encabezados
            writer.writerow(['Vendedor', 'Total_Citas', 'Horarios', 'Compradores', 'Preferencias_Cumplidas'])
            
            # Datos
            for vendedor, datos in resultado["resumen_por_vendedor"].items():
                # Calcular preferencias cumplidas para este vendedor
                preferencias_vendedor = resultado.get("preferencias_cumplidas", {}).get(vendedor, [])
                cumplidas = sum(1 for pref in preferencias_vendedor if pref.get("cumplida", False))
                total_pref = len(preferencias_vendedor)
                preferencia_cumplida = f"{cumplidas}/{total_pref}" if total_pref > 0 else "0/0"
                
                if datos['total_citas'] > 0:
                    horarios = [cita['horario'] for cita in datos['citas']]
                    compradores = [cita['comprador'] for cita in datos['citas']]
                    writer.writerow([
                        vendedor, 
                        datos['total_citas'],
                        '; '.join(horarios),
                        '; '.join(compradores),
                        preferencia_cumplida
                    ])
                else:
                    writer.writerow([vendedor, 0, 'Sin citas', 'Sin citas', preferencia_cumplida])
        
        print(f"Resumen de vendedores exportado a: {ruta_archivo}")

    def generar_excel_completo(self, resultado: Dict, nombre_archivo: str = "agenda_rueda_negocios.xlsx"):
        """Genera un archivo Excel completo con múltiples hojas basado en agenda_completa.json"""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            print("Error: Se requiere instalar openpyxl para generar Excel.")
            print("Ejecuta: pip install openpyxl")
            return None
        
        ruta_archivo = f"c:\\Users\\angel\\OneDrive\\Escritorio\\escritorio\\Angela\\Programación agenda\\{nombre_archivo}"
        
        # Crear workbook y eliminar hoja por defecto
        wb = openpyxl.Workbook()
        wb.remove(wb.active)
        
        # Estilos
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
        subheader_fill = PatternFill(start_color="3498DB", end_color="3498DB", fill_type="solid")
        accent_fill = PatternFill(start_color="ECF0F1", end_color="ECF0F1", fill_type="solid")
        border = Border(left=Side(style='thin'), right=Side(style='thin'), 
                       top=Side(style='thin'), bottom=Side(style='thin'))
        center_align = Alignment(horizontal='center', vertical='center')
        
        # 1. HOJA: Resumen General
        ws_resumen = wb.create_sheet("Resumen General")
        
        # Título
        ws_resumen['A1'] = "AGENDA RUEDA DE NEGOCIOS - RESUMEN GENERAL"
        ws_resumen['A1'].font = Font(bold=True, size=16)
        ws_resumen.merge_cells('A1:D1')
        
        # Configuración
        config = resultado["configuracion"]
        stats = resultado["estadisticas"]
        
        ws_resumen['A3'] = "CONFIGURACIÓN"
        ws_resumen['A3'].font = header_font
        ws_resumen['A3'].fill = header_fill
        
        config_data = [
            ["Vendedores", config['vendedores']],
            ["Compradores", config['compradores']],
            ["Duración total", config['duracion_total']],
            ["Duración por cita", config['duracion_cita']],
            ["Horario", f"{config['horario_inicio']} - {config['horario_fin']}"],
            ["Slots disponibles", config['slots_disponibles']],
            ["Máx. citas por vendedor", config['max_citas_por_vendedor']],
            ["Vendedores por cita", config['vendedores_por_cita']],
            ["Preferencias cargadas", config['preferencias_cargadas']]
        ]
        
        for i, (label, value) in enumerate(config_data, 4):
            ws_resumen[f'A{i}'] = label
            ws_resumen[f'B{i}'] = value
            ws_resumen[f'A{i}'].fill = accent_fill
        
        # Estadísticas
        ws_resumen['A14'] = "ESTADÍSTICAS"
        ws_resumen['A14'].font = header_font
        ws_resumen['A14'].fill = header_fill
        
        stats_data = [
            ["Total citas programadas", stats['total_citas_programadas']],
            ["Total encuentros individuales", stats['total_encuentros_individuales']],
            ["Utilización de slots", stats['porcentaje_utilizacion_slots']],
            ["Promedio citas por vendedor", stats['citas_promedio_por_vendedor']],
            ["Promedio citas por comprador", stats['citas_promedio_por_comprador']],
            ["Preferencias cumplidas", f"{stats['preferencias_cumplidas']}/{stats['total_preferencias']}"],
            ["% Cumplimiento preferencias", f"{(stats['preferencias_cumplidas']/stats['total_preferencias']*100):.1f}%"]
        ]
        
        for i, (label, value) in enumerate(stats_data, 15):
            ws_resumen[f'A{i}'] = label
            ws_resumen[f'B{i}'] = value
            ws_resumen[f'A{i}'].fill = accent_fill
        
        # 2. HOJA: Agenda por Horarios
        ws_agenda = wb.create_sheet("Agenda por Horarios")
        
        # Encabezados
        headers = ['Horario', 'Cita #', 'Comprador', 'Vendedor 1', 'Vendedor 2', 'Vendedor 3']
        for col, header in enumerate(headers, 1):
            cell = ws_agenda.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = border
        
        row = 2
        for horario, citas in resultado["agenda"].items():
            for i, cita in enumerate(citas, 1):
                ws_agenda.cell(row=row, column=1, value=horario)
                ws_agenda.cell(row=row, column=2, value=f"Cita {i}")
                ws_agenda.cell(row=row, column=3, value=cita['comprador'])
                
                # Vendedores (máximo 3)
                for j, vendedor in enumerate(cita['vendedores'][:3], 4):
                    ws_agenda.cell(row=row, column=j, value=vendedor)
                
                # Aplicar estilos
                for col in range(1, 7):
                    cell = ws_agenda.cell(row=row, column=col)
                    cell.border = border
                    if col == 3:  # Comprador
                        cell.fill = PatternFill(start_color="E74C3C", end_color="E74C3C", fill_type="solid")
                        cell.font = Font(color="FFFFFF", bold=True)
                    elif col >= 4:  # Vendedores
                        cell.fill = PatternFill(start_color="27AE60", end_color="27AE60", fill_type="solid")
                        cell.font = Font(color="FFFFFF")
                
                row += 1
        
        # 3. HOJA: Resumen por Vendedores
        ws_vendedores = wb.create_sheet("Resumen Vendedores")
        
        headers = ['Vendedor', 'Total Citas', 'Preferencias Cumplidas', 'Lista de Compradores', 'Horarios']
        for col, header in enumerate(headers, 1):
            cell = ws_vendedores.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = border
        
        row = 2
        for vendedor, datos in resultado["resumen_por_vendedor"].items():
            # Calcular preferencias cumplidas
            preferencias_vendedor = resultado.get("preferencias_cumplidas", {}).get(vendedor, [])
            cumplidas = sum(1 for pref in preferencias_vendedor if pref.get("cumplida", False))
            total_pref = len(preferencias_vendedor)
            
            ws_vendedores.cell(row=row, column=1, value=vendedor)
            ws_vendedores.cell(row=row, column=2, value=datos['total_citas'])
            ws_vendedores.cell(row=row, column=3, value=f"{cumplidas}/{total_pref}")
            
            if datos['total_citas'] > 0:
                compradores = [cita['comprador'] for cita in datos['citas']]
                horarios = [cita['horario'] for cita in datos['citas']]
                ws_vendedores.cell(row=row, column=4, value='; '.join(compradores))
                ws_vendedores.cell(row=row, column=5, value='; '.join(horarios))
            else:
                ws_vendedores.cell(row=row, column=4, value="Sin citas")
                ws_vendedores.cell(row=row, column=5, value="Sin citas")
            
            # Aplicar estilos
            for col in range(1, 6):
                cell = ws_vendedores.cell(row=row, column=col)
                cell.border = border
                if datos['total_citas'] == 0:
                    cell.fill = PatternFill(start_color="F39C12", end_color="F39C12", fill_type="solid")
            
            row += 1
        
        # 4. HOJA: Resumen por Compradores
        ws_compradores = wb.create_sheet("Resumen Compradores")
        
        headers = ['Comprador', 'Total Citas', 'Lista de Vendedores', 'Horarios']
        for col, header in enumerate(headers, 1):
            cell = ws_compradores.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = border
        
        row = 2
        for comprador, datos in resultado["resumen_por_comprador"].items():
            ws_compradores.cell(row=row, column=1, value=comprador)
            ws_compradores.cell(row=row, column=2, value=datos['total_citas'])
            
            if datos['total_citas'] > 0:
                vendedores = []
                horarios = [cita['horario'] for cita in datos['citas']]
                for cita in datos['citas']:
                    vendedores.extend(cita['vendedores'])
                
                ws_compradores.cell(row=row, column=3, value='; '.join(vendedores))
                ws_compradores.cell(row=row, column=4, value='; '.join(horarios))
            else:
                ws_compradores.cell(row=row, column=3, value="Sin citas")
                ws_compradores.cell(row=row, column=4, value="Sin citas")
            
            # Aplicar estilos
            for col in range(1, 5):
                cell = ws_compradores.cell(row=row, column=col)
                cell.border = border
            
            row += 1
        
        # 5. HOJA: Matriz Compradores-Horarios
        ws_matriz = wb.create_sheet("Matriz Compradores-Horarios")
        
        # Crear lista de todos los compradores únicos
        compradores_unicos = list(resultado["resumen_por_comprador"].keys())
        
        # Crear lista de todos los horarios de la agenda
        horarios_agenda = [horario for horario in resultado["agenda"].keys()]
        
        # Encabezado: primera celda vacía, luego compradores
        ws_matriz.cell(row=1, column=1, value="FRANJA HORARIA")
        ws_matriz.cell(row=1, column=1).font = header_font
        ws_matriz.cell(row=1, column=1).fill = header_fill
        ws_matriz.cell(row=1, column=1).alignment = center_align
        ws_matriz.cell(row=1, column=1).border = border
        
        # Agregar compradores como encabezados de columnas
        for col, comprador in enumerate(compradores_unicos, 2):
            cell = ws_matriz.cell(row=1, column=col, value=comprador)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = border
        
        # Llenar la matriz
        for row_idx, horario in enumerate(horarios_agenda, 2):
            # Primera columna: horario
            cell_horario = ws_matriz.cell(row=row_idx, column=1, value=horario)
            cell_horario.font = Font(bold=True)
            cell_horario.fill = subheader_fill
            cell_horario.alignment = center_align
            cell_horario.border = border
            
            # Para cada comprador, buscar los vendedores en este horario
            for col_idx, comprador in enumerate(compradores_unicos, 2):
                vendedores_en_horario = []
                
                # Buscar citas de este comprador en este horario
                if horario in resultado["agenda"]:
                    for cita in resultado["agenda"][horario]:
                        if cita["comprador"] == comprador:
                            vendedores_en_horario.extend(cita["vendedores"])
                
                # Escribir vendedores en la celda (separados por comas si hay múltiples)
                vendedores_texto = ", ".join(vendedores_en_horario) if vendedores_en_horario else ""
                cell = ws_matriz.cell(row=row_idx, column=col_idx, value=vendedores_texto)
                cell.border = border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                
                # Colorear según si hay citas o no
                if vendedores_en_horario:
                    cell.fill = PatternFill(start_color="E8F5E8", end_color="E8F5E8", fill_type="solid")
                    cell.font = Font(bold=True)
                else:
                    cell.fill = PatternFill(start_color="F8F8F8", end_color="F8F8F8", fill_type="solid")
        
        # 6. HOJA: Preferencias Detalladas
        ws_preferencias = wb.create_sheet("Preferencias Detalladas")
        
        headers = ['Vendedor', 'Comprador Preferido', 'Estado', 'Cumplida']
        for col, header in enumerate(headers, 1):
            cell = ws_preferencias.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = border
        
        row = 2
        for vendedor, preferencias_vendedor in resultado["preferencias_cumplidas"].items():
            for pref in preferencias_vendedor:
                ws_preferencias.cell(row=row, column=1, value=vendedor)
                ws_preferencias.cell(row=row, column=2, value=pref['comprador'])
                ws_preferencias.cell(row=row, column=3, value="✓ Cumplida" if pref['cumplida'] else "✗ No cumplida")
                ws_preferencias.cell(row=row, column=4, value="SÍ" if pref['cumplida'] else "NO")
                
                # Aplicar estilos
                for col in range(1, 5):
                    cell = ws_preferencias.cell(row=row, column=col)
                    cell.border = border
                    if pref['cumplida']:
                        cell.fill = PatternFill(start_color="D5F4E6", end_color="D5F4E6", fill_type="solid")
                    else:
                        cell.fill = PatternFill(start_color="FDF2F2", end_color="FDF2F2", fill_type="solid")
                
                row += 1
        
        # Ajustar ancho de columnas para todas las hojas
        for ws in wb.worksheets:
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
        
        # Guardar archivo
        wb.save(ruta_archivo)
        
        print(f"\n✓ Archivo Excel generado: {ruta_archivo}")
        print(f"  - 6 hojas: Resumen General, Agenda por Horarios, Resumen Vendedores, Resumen Compradores, Matriz Compradores-Horarios, Preferencias Detalladas")
        print(f"  - Formato profesional con colores y estilos")
        print(f"  - {stats['total_citas_programadas']} citas y {stats['total_encuentros_individuales']} encuentros detallados")
        print(f"  - {stats['preferencias_cumplidas']}/{stats['total_preferencias']} preferencias analizadas")
        print(f"\n📊 MATRIZ MAESTRA: Hoja 'Matriz Compradores-Horarios' - Grilla completa con compradores como filas, horarios como columnas y vendedores en cada celda")
        
        return ruta_archivo

    def crear_archivo_ejemplo_preferencias(self, nombre_archivo: str = "ejemplo_preferencias.csv"):
        """Crea un archivo de ejemplo para mostrar el formato de preferencias"""
        ruta_archivo = f"c:\\Users\\angel\\OneDrive\\Escritorio\\escritorio\\Angela\\Programación agenda\\{nombre_archivo}"
        
        with open(ruta_archivo, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.writer(csvfile)
            
            # Encabezados
            writer.writerow(['Nombre_Vendedor', 'Comprador_Preferido'])
            
            # Ejemplos con nombres más realistas
            ejemplos = [
                ['Café Del Tajo', 'BOX BRAND'],
                ['Café Del Tajo', 'PROCOLOMBIA'],  # Mismo vendedor, múltiples compradores
                ['Café Del Tajo', 'ENCADENAMIENTOS PRODUCTIVOS'],
                ['Mujeres cafeteras ALMA VERDE', 'CAFÉ MOLINA'],
                ['Mujeres cafeteras ALMA VERDE', 'COLFRESH COFFEE'],
                ['Cervecería Montelargo', 'BOX BRAND'],
                ['Cervecería Montelargo', 'NEIRA YORK COFFEE'],
                ['Productora Agrícola XYZ', 'INMERSSO BOUTIQUE'],
                ['Cooperativa Cafetera ABC', 'CAFÉ MOLINA'],
                ['Empresa Sostenible 123', 'PROCOLOMBIA'],
                ['Vendedor Local DEF', 'COLFRESH COFFEE'],
                ['Microempresa GHI', 'BOX BRAND'],
            ]
            
            for ejemplo in ejemplos:
                writer.writerow(ejemplo)
        
        print(f"Archivo de ejemplo creado: {ruta_archivo}")
        print("NOTA: Ahora cada vendedor puede tener MÚLTIPLES compradores preferidos.")
        print("      Simplemente repite el nombre del vendedor en varias filas.")
        print("      ¡Usa nombres REALES! El programa los detectará automáticamente.")
        print("Edita este archivo con tus preferencias reales y úsalo con cargar_preferencias_archivo()")
        return ruta_archivo

    def generar_documentos_word_vendedores(self, resultado: Dict):
        """Genera un documento Word individual para cada vendedor con sus citas"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            print("Error: Se requiere instalar python-docx para generar documentos Word.")
            print("Ejecuta: pip install python-docx")
            return

        # Crear carpeta para documentos si no existe
        carpeta_docs = "c:\\Users\\angel\\OneDrive\\Escritorio\\escritorio\\Angela\\Programación agenda\\documentos_vendedores"
        import os
        if not os.path.exists(carpeta_docs):
            os.makedirs(carpeta_docs)

        vendedores_procesados = 0
        
        # Generar un documento para cada vendedor
        for vendedor, datos_vendedor in resultado["resumen_por_vendedor"].items():
            if datos_vendedor["total_citas"] > 0:  # Solo generar para vendedores con citas
                doc = Document()
                
                # Título principal: nombre del vendedor
                titulo = doc.add_heading(vendedor, 0)
                titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Agregar espacio
                doc.add_paragraph("")
                
                # Información general
                info_p = doc.add_paragraph()
                info_p.add_run("Total de citas programadas: ").bold = True
                info_p.add_run(str(datos_vendedor["total_citas"]))
                
                # Agregar espacio
                doc.add_paragraph("")
                
                # Crear tabla con citas
                # Encabezados: Franja Horaria | Mesa | Comprador
                tabla = doc.add_table(rows=1, cols=3)
                tabla.style = 'Table Grid'
                tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Encabezados de la tabla
                encabezados = tabla.rows[0].cells
                encabezados[0].text = "FRANJA HORARIA"
                encabezados[1].text = "MESA"
                encabezados[2].text = "COMPRADOR"
                
                # Formatear encabezados
                for celda in encabezados:
                    for parrafo in celda.paragraphs:
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in parrafo.runs:
                            run.font.bold = True
                
                # Agregar las citas del vendedor
                for cita in datos_vendedor["citas"]:
                    fila = tabla.add_row().cells
                    fila[0].text = cita["horario"]
                    fila[1].text = ""  # Mesa vacía como solicitaste
                    fila[2].text = cita["comprador"]
                    
                    # Centrar el contenido de las celdas
                    for celda in fila:
                        for parrafo in celda.paragraphs:
                            parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Ajustar ancho de columnas
                for fila in tabla.rows:
                    fila.cells[0].width = Inches(2.5)  # Franja horaria
                    fila.cells[1].width = Inches(1.5)  # Mesa
                    fila.cells[2].width = Inches(3.0)  # Comprador
                
                # Agregar espacio al final
                doc.add_paragraph("")
                
                # Nota al pie
                nota = doc.add_paragraph()
                nota.add_run("Nota: ").bold = True
                nota.add_run("Complete la columna 'Mesa' según la asignación del evento.")
                nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Guardar el documento
                # Limpiar caracteres problemáticos del nombre del archivo
                caracteres_problematicos = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
                nombre_limpio = vendedor
                for char in caracteres_problematicos:
                    nombre_limpio = nombre_limpio.replace(char, '_')
                
                nombre_archivo = f"{nombre_limpio}.docx"
                ruta_completa = os.path.join(carpeta_docs, nombre_archivo)
                
                try:
                    doc.save(ruta_completa)
                    vendedores_procesados += 1
                    print(f"✓ Documento generado: {nombre_archivo}")
                except Exception as e:
                    print(f"✗ Error al guardar documento para {vendedor}: {e}")
        
        print(f"\n📄 DOCUMENTOS WORD GENERADOS:")
        print(f"   • Total documentos: {vendedores_procesados}")
        print(f"   • Carpeta: {carpeta_docs}")
        print(f"   • Cada documento contiene: Nombre del vendedor, tabla con franja horaria, mesa (vacía) y comprador")
        
        return vendedores_procesados


def main():
    """Función principal"""
    print("Iniciando programa de agenda para rueda de negocios...")
    print("Nuevo formato: 3 vendedores por cita con 1 comprador")
    
    # Crear instancia del organizador
    organizador = AgendaRuedaNegocios()
    
    # Crear archivo de ejemplo para preferencias
    print("\n1. Creando archivo de ejemplo para preferencias...")
    archivo_ejemplo = organizador.crear_archivo_ejemplo_preferencias()
    
    # Intentar cargar preferencias si existe un archivo
    archivo_preferencias = "c:\\Users\\angel\\OneDrive\\Escritorio\\escritorio\\Angela\\Programación agenda\\preferencias_multiples.csv"
    if os.path.exists(archivo_preferencias):
        print(f"\n2. Cargando preferencias desde {archivo_preferencias}...")
        organizador.cargar_preferencias_archivo(archivo_preferencias)
    else:
        archivo_preferencias_alt = "c:\\Users\\angel\\OneDrive\\Escritorio\\escritorio\\Angela\\Programación agenda\\preferencias_citas.csv"
        if os.path.exists(archivo_preferencias_alt):
            print(f"\n2. Cargando preferencias desde {archivo_preferencias_alt}...")
            organizador.cargar_preferencias_archivo(archivo_preferencias_alt)
        else:
            print(f"\n2. No se encontró archivo de preferencias")
            print("   Puedes crear uno basado en el ejemplo generado.")
    
    # Generar la agenda
    print("\n3. Generando agenda optimizada...")
    resultado = organizador.generar_agenda_optimizada()
    
    # Mostrar resultados
    print("\n4. Mostrando resultados...")
    organizador.imprimir_agenda(resultado)
    
    # Exportar archivos
    print("\n5. Exportando archivos...")
    organizador.exportar_a_csv(resultado)
    organizador.exportar_resumen_vendedores(resultado)
    
    # Generar archivo Excel completo
    print("\n6. Generando archivo Excel...")
    organizador.generar_excel_completo(resultado)
    
    # Generar documentos Word para vendedores
    print("\n7. Generando documentos Word individuales para vendedores...")
    organizador.generar_documentos_word_vendedores(resultado)
    
    # Guardar resultado completo en JSON
    with open("c:\\Users\\angel\\OneDrive\\Escritorio\\escritorio\\Angela\\Programación agenda\\agenda_completa.json", 'w', encoding='utf-8') as f:
        json.dump(resultado, f, indent=2, ensure_ascii=False)
    
    print(f"\nResultado completo guardado en: agenda_completa.json")
    print("\nPrograma finalizado exitosamente!")
    print("\nPARA USAR TUS PROPIAS PREFERENCIAS:")
    print("1. Edita el archivo 'ejemplo_preferencias.csv' con tus datos reales")
    print("2. Cada vendedor puede aparecer en MÚLTIPLES filas con diferentes compradores")
    print("3. ¡USA NOMBRES REALES! El programa detectará automáticamente todos los participantes")
    print("4. Guárdalo como 'preferencias_multiples.csv' o 'preferencias_citas.csv'")
    print("5. Ejecuta el programa nuevamente")
    print("\nEjemplo del nuevo formato:")
    print("Café Del Tajo,BOX BRAND")
    print("Café Del Tajo,PROCOLOMBIA  <- Mismo vendedor, otro comprador")
    print("Mujeres cafeteras ALMA VERDE,CAFÉ MOLINA  <- Vendedor diferente")


if __name__ == "__main__":
    main()
