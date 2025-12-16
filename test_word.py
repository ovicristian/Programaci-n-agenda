#!/usr/bin/env python3
"""Test para generar documentos Word individuales"""

import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def generar_documentos_word_vendedores_test():
    """Función de prueba para generar documentos Word"""
    
    # Cargar datos desde el JSON generado
    with open("agenda_completa.json", 'r', encoding='utf-8') as f:
        resultado = json.load(f)
    
    # Crear carpeta para documentos si no existe
    carpeta_docs = "documentos_vendedores"
    if not os.path.exists(carpeta_docs):
        os.makedirs(carpeta_docs)
        print(f"Carpeta creada: {carpeta_docs}")

    vendedores_procesados = 0
    
    # Generar un documento para cada vendedor
    for vendedor, datos_vendedor in resultado["resumen_por_vendedor"].items():
        if datos_vendedor["total_citas"] > 0:  # Solo generar para vendedores con citas
            print(f"Procesando vendedor: {vendedor}")
            
            doc = Document()
            
            # Título principal: nombre del vendedor
            titulo = doc.add_heading(vendedor, 0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Agregar espacio
            doc.add_paragraph("")
            
            # Información general
            info_p = doc.add_paragraph()
            info_p.add_run("Total de citas programadas: ").bold = True
            info_p.add_run(str(datos_vendedor["total_citas"]))
            
            # Agregar espacio
            doc.add_paragraph("")
            
            # Crear tabla con citas
            # Encabezados: Franja Horaria | Mesa | Comprador
            tabla = doc.add_table(rows=1, cols=3)
            tabla.style = 'Table Grid'
            tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Encabezados de la tabla
            encabezados = tabla.rows[0].cells
            encabezados[0].text = "FRANJA HORARIA"
            encabezados[1].text = "MESA"
            encabezados[2].text = "COMPRADOR"
            
            # Formatear encabezados
            for celda in encabezados:
                for parrafo in celda.paragraphs:
                    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in parrafo.runs:
                        run.font.bold = True
            
            # Agregar las citas del vendedor
            for cita in datos_vendedor["citas"]:
                fila = tabla.add_row().cells
                fila[0].text = cita["horario"]
                fila[1].text = ""  # Mesa vacía como solicitaste
                fila[2].text = cita["comprador"]
                
                # Centrar el contenido de las celdas
                for celda in fila:
                    for parrafo in celda.paragraphs:
                        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Ajustar ancho de columnas
            for fila in tabla.rows:
                fila.cells[0].width = Inches(2.5)  # Franja horaria
                fila.cells[1].width = Inches(1.5)  # Mesa
                fila.cells[2].width = Inches(3.0)  # Comprador
            
            # Agregar espacio al final
            doc.add_paragraph("")
            
            # Nota al pie
            nota = doc.add_paragraph()
            nota.add_run("Nota: ").bold = True
            nota.add_run("Complete la columna 'Mesa' según la asignación del evento.")
            nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Guardar el documento
            # Limpiar caracteres problemáticos del nombre del archivo
            caracteres_problematicos = ['/', '\\', ':', '*', '?', '"', '<', '>', '|']
            nombre_limpio = vendedor
            for char in caracteres_problematicos:
                nombre_limpio = nombre_limpio.replace(char, '_')
            
            nombre_archivo = f"{nombre_limpio}.docx"
            ruta_completa = os.path.join(carpeta_docs, nombre_archivo)
            
            try:
                doc.save(ruta_completa)
                vendedores_procesados += 1
                print(f"✓ Documento generado: {nombre_archivo}")
            except Exception as e:
                print(f"✗ Error al guardar documento para {vendedor}: {e}")
    
    print(f"\n📄 DOCUMENTOS WORD GENERADOS:")
    print(f"   • Total documentos: {vendedores_procesados}")
    print(f"   • Carpeta: {carpeta_docs}")
    print(f"   • Cada documento contiene: Nombre del vendedor, tabla con franja horaria, mesa (vacía) y comprador")
    
    return vendedores_procesados

if __name__ == "__main__":
    print("🚀 Iniciando generación de documentos Word para vendedores...")
    generar_documentos_word_vendedores_test()
    print("✅ Proceso completado!")
